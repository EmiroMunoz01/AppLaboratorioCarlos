# imprimir_cotizacion.py
import os
from docx import Document
# Ya no necesitamos importar Decimal
from modelo import listar_items_por_cotizacion, obtener_vehiculo_por_placa

# En tu archivo: imprimir_cotizacion.py

import os
from docx import Document
from modelo import listar_items_por_cotizacion, obtener_vehiculo_por_placa

def generar_cotizacion_para_placa(
    placa: str,
    fecha: str,  # La fecha llega formateada, ej: '16-10-2025'
    cot_id: int,
    output_dir: str = r"C:\Users\Admin\Desktop\AppLaboratorioCarlos\docs\COTIZACIONES",
) -> str:
    """
    Genera un documento .docx de cotización y reemplaza el marcador de fecha
    en párrafos y tablas.
    """
    # 1. Cargar la plantilla de Word
    plantilla_path = os.path.join(os.path.dirname(__file__), "plantillas", "cotizacion.docx")
    doc = Document(plantilla_path)

    # --- INICIO DEL BLOQUE DE REEMPLAZO MEJORADO ---
    # Marcador que vamos a buscar en la plantilla
    marcador_fecha = "Fecha de emisión" 
    
    # Reemplazar en todos los párrafos del documento
    for p in doc.paragraphs:
        if marcador_fecha in p.text:
            # Limpiamos el párrafo y lo reescribimos para evitar duplicados
            p.text = "" 
            p.add_run(f"{marcador_fecha}: {fecha}")

    # Reemplazar en todas las tablas del documento
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if marcador_fecha in cell.text:
                    # Limpiamos la celda y la reescribimos
                    cell.text = ""
                    cell.add_paragraph(f"{marcador_fecha}: {fecha}")
    # --- FIN DEL BLOQUE DE REEMPLAZO MEJORADO ---

    # 2. Obtener los datos restantes de la base de datos
    datos_vehiculo = obtener_vehiculo_por_placa(placa)
    items_cotizacion = listar_items_por_cotizacion(cot_id)

    # 3. Rellenar el resto de la información (clientes, ítems, total)
    # ... (Aquí va tu código existente para rellenar las tablas de cliente e ítems)
    # ... (Asegúrate de que esta lógica siga funcionando como antes)

    # Por ejemplo, el código para rellenar los items y el total:
    tbl_items = doc.tables[1]  # Asumiendo que es la segunda tabla
    suma_total = 0
    for item in items_cotizacion:
        celdas_fila_nueva = tbl_items.add_row().cells
        celdas_fila_nueva[0].text = item.get('articulo') or ""
        celdas_fila_nueva[1].text = item.get('descripcion', "")
        celdas_fila_nueva[2].text = f"{item.get('cantidad', 0):,g}"
        celdas_fila_nueva[3].text = f"{item.get('precio_unit', 0):,}"
        celdas_fila_nueva[4].text = f"{item.get('total', 0):,}"
        suma_total += item.get('total', 0)

    for p in doc.paragraphs:
        if "VALOR FINAL COTIZACION" in p.text:
            p.text = "" 
            p.add_run(f"VALOR FINAL: ${suma_total:,}").bold = True


    # 4. Guardar el documento final
    os.makedirs(output_dir, exist_ok=True)
    nombre_archivo = f"cotizacion_{placa}_{cot_id}.docx"
    ruta_guardado = os.path.join(output_dir, nombre_archivo)
    doc.save(ruta_guardado)

    return ruta_guardado
